from datetime import time

import pandas as pd
import pymysql
import mysql.connector  # ✅ Correct way to import
import openai
import json
from flask import Flask, render_template, request, jsonify, session, url_for
from scipy.stats import stats

pymysql.install_as_MySQLdb()
import os
import re
import uuid
import arabic_reshaper
import nltk
import json
from flask import Flask, render_template, request, redirect, url_for, session
from flask_mysqldb import MySQL
from matplotlib import rcParams
from sympy.physics.control.control_plots import matplotlib
from werkzeug.security import generate_password_hash, check_password_hash
from collections import Counter
import matplotlib.pyplot as plt
import arabic_reshaper
from bidi.algorithm import get_display
from nltk.util import ngrams
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from wordcloud import WordCloud
import MySQLdb.cursors
import mishkal
from camel_tools.utils.charmap import CharMapper
from camel_tools.tokenizers.word import simple_word_tokenize
# from camel_tools.stem.isri import ISRIStemmer
from nltk.stem.isri import ISRIStemmer
import smtplib
from email.mime.text import MIMEText
# New imports for NER and Keyword Extraction
from transformers import AutoTokenizer, AutoModelForTokenClassification, AutoModel
import torch
import numpy as np
from keybert import KeyBERT
import uuid
from docx import Document

from flask import Response, stream_with_context

# pip install python-docx

# Existing imports
from flask import Flask, render_template, request, redirect, url_for, session
from flask_mysqldb import MySQL
# ... (rest of your existing imports)

# Add the new import here
from werkzeug.utils import secure_filename

from flask import Flask, render_template, request, redirect, url_for, session
# Remove this unused import
# from datetime import time
from io import StringIO, BytesIO
import csv
import xlsxwriter

app = Flask(__name__)


# Define utility functions BEFORE routes
def read_file_content(file_path):
    """Read content from either .txt or .docx files."""
    try:
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None

        extension = os.path.splitext(file_path)[1].lower()

        if extension == '.txt':
            encodings = ['utf-8', 'utf-8-sig', 'cp1256', 'windows-1256', 'utf-16', 'utf-16le', 'utf-16be']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        if content.strip():
                            return content
                except UnicodeDecodeError:
                    continue
            return None

        elif extension == '.docx':
            try:
                doc = Document(file_path)
                content = []

                # Get text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append(para.text)

                # Get text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            content.append(" | ".join(row_text))

                return "\n".join(content) if content else None

            except Exception as e:
                print(f"Error reading docx file: {e}")
                return None

        else:
            print(f"Unsupported file extension: {extension}")
            return None

    except Exception as e:
        print(f"Error reading file: {e}")
        return None


# Your Flask app initialization

# ... (rest of your existing code)


app.secret_key = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'uploads'

app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'root'  # Default MAMP password (if unchanged)
app.config['MYSQL_DB'] = 'bayyin'
app.config['MYSQL_SSL_DISABLED'] = True  # Disable SSL

# Initialize MySQL
mysql = MySQL(app)

matplotlib.use('Agg')  # Ensure matplotlib doesn't require an active display

# Set up Arabic text rendering in matplotlib
rcParams['font.family'] = 'sans-serif'
rcParams['font.sans-serif'] = ['Amiri']  # Ensure 'Amiri' font is available
rcParams['axes.unicode_minus'] = False

# Download required NLTK packages
nltk.download('punkt')
nltk.download('stopwords')

# Use NLTK's Arabic stopwords list
arabic_stopwords = set(nltk.corpus.stopwords.words("arabic"))

# Manually add missing entries to the stopword list
additional_stopwords = {
    "او", "ان", "انها", "أنه", "انه", "أن", "أنها", "إلى", "الي", "على", "علي",
    "وهي", "ولم", "وان", "الا", "بهذه", "ومع", "ام", "لدي", "فلم", "وانها",
    "فانهم", "وهناك", "لهذه", "وذلك", "وعلي", "اذ", "فان", "لان", "وحتي", "وفي",
    "ولان", "اذا", "وهذا", "وبهذا", "فهذه", "كانت"
}
arabic_stopwords.update(additional_stopwords)

# Initialize NLTK stemmer for Arabic
stemmer = ISRIStemmer()

# Initialize models for NER and Keyword Extraction
model_ner_name = "marefa-nlp/marefa-ner"
tokenizer_ner = AutoTokenizer.from_pretrained(model_ner_name)
model_ner = AutoModelForTokenClassification.from_pretrained(model_ner_name)

model_kw_name = "aubmindlab/bert-base-arabertv02"
tokenizer_kw = AutoTokenizer.from_pretrained(model_kw_name)
model_kw = AutoModel.from_pretrained(model_kw_name)
kw_model = KeyBERT(model_kw)

# Custom labels for NER
custom_labels = [
    "O", "B-job", "I-job", "B-nationality", "B-person", "I-person",
    "B-location", "B-time", "I-time", "B-event", "I-event",
    "B-organization", "I-organization", "I-location", "I-nationality",
    "B-product", "I-product", "B-artwork", "I-artwork"
]

# Path to save uploaded files
# UPLOAD_FOLDER = 'uploads/'
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)  # Ensure directory exists

# At the start of your app configuration, before any routes
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'uploads')

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER, mode=0o755)

app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize character mapper for normalization
mapper = CharMapper.builtin_mapper('arclean')

# List of common Arabic words with valid double letters (without شدّة)
valid_double_letter_words = [
    "مرر", "سدد", "ضرر", "حلل", "قرر", "جرر", "كرر", "فرر", "مدد", "ردد",
    "شدد", "هدد", "خطط", "برر", "عمم", "رمم", "حلل", "سدد", "مدد", "كرّر",
    "هدد", "ضرر", "مرر", "قلل", "خطط", "مدد", "برر", "رمم", "زرر", "جفف",
    "تتلألأ", "اللغه", "اللغة", "تتلألأ"
]

smtp_server = "smtp.gmail.com"
smtp_port = 587


@app.route('/')
def root():
    if 'loggedin' in session:
        return redirect(url_for('upload_file'))
    return redirect(url_for('login'))


# Authentication and home page routes
@app.route('/index')
def index():
    if 'loggedin' in session:
        return redirect(url_for('upload_file'))
    return redirect(url_for('login'))


from werkzeug.security import check_password_hash


@app.route('/auth', methods=['POST'])
def auth():
    action = request.form.get('action')
    if action == 'login':
        # Get login credentials from form
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()

        if not username or not password:
            return render_template('login.html', error_message="يرجى إدخال اسم المستخدم وكلمة المرور.")

        # Query the user from the database
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM tbl_users WHERE username = %s', (username,))
        user = cursor.fetchone()

        if user:
            # Check if the user has verified their email
            if not user['is_verified']:
                return render_template('login.html', error_message="يرجى التحقق من بريدك الإلكتروني قبل تسجيل الدخول.")

            # Verify the password
            if check_password_hash(user['password'], password):
                session['loggedin'] = True
                session['username'] = user['username']
                session['id'] = user['id']
                return redirect(url_for('home'))
            else:
                return render_template('login.html', error_message="اسم المستخدم أو كلمة المرور غير صحيحة!")
        else:
            return render_template('login.html', error_message="اسم المستخدم أو كلمة المرور غير صحيحة!")

    elif action == 'register':
        # Get registration details from form
        first_name = request.form.get('first_name', '').strip()
        second_name = request.form.get('second_name', '').strip()
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()

        # Validate inputs
        if not all([first_name, second_name, username, email, password]):
            return render_template('login.html', error_message="يرجى ملء جميع الحقول المطلوبة.")

        if not first_name.isalpha() or not second_name.isalpha():
            return render_template('login.html', error_message="يجب ألا يحتوي الاسم الأول أو الثاني على أرقام.")

        if "@" not in email or "." not in email:
            return render_template('login.html', error_message="يرجى إدخال بريد إلكتروني صحيح.")

        # Check if the username is already taken
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT * FROM tbl_users WHERE username = %s', (username,))
        existing_user = cursor.fetchone()

        if existing_user:
            return render_template('login.html', error_message="اسم المستخدم موجود بالفعل. يرجى اختيار اسم مستخدم آخر.")

        # Hash the password
        password_hashed = generate_password_hash(password)
        verification_token = str(uuid.uuid4())

        try:
            # Insert the new user into the database with verification status as false
            cursor.execute(
                'INSERT INTO tbl_users (first_name, second_name, username, email, password, verification_token, is_verified) VALUES (%s, %s, %s, %s, %s, %s, %s)',
                (first_name, second_name, username, email, password_hashed, verification_token, False)
            )
            mysql.connection.commit()

            # Send a verification email
            send_verification_email(email, verification_token)
            return render_template('verify_email.html',
                                   message="تم إرسال بريد إلكتروني للتحقق. يرجى التحقق من بريدك الإلكتروني.")
        except Exception as e:
            print(f"Error during registration: {e}")
            return f"Registration failed: {str(e)}"

    else:
        return render_template('login.html', error_message="Invalid action")


from flask import redirect, url_for


@app.route('/verify_email/<token>')
def verify_email(token):
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cursor.execute('SELECT * FROM tbl_users WHERE verification_token = %s', (token,))
    user = cursor.fetchone()

    if user and not user['is_verified']:
        cursor.execute('UPDATE tbl_users SET is_verified = %s WHERE id = %s', (True, user['id']))
        mysql.connection.commit()
        return redirect(url_for('login', success_message="تم التحقق من بريدك الإلكتروني. يمكنك الآن تسجيل الدخول."))
    elif user and user['is_verified']:
        return redirect(url_for('login', error_message="هذا البريد الإلكتروني تم التحقق منه مسبقاً."))
    else:
        return redirect(url_for('login', error_message="رابط التحقق غير صالح."))


@app.route('/login')
def login():
    success_message = request.args.get('success_message', '')
    error_message = request.args.get('error_message', '')
    return render_template('login.html', success_message=success_message, error_message=error_message)


def send_verification_email(email, verification_token):
    verification_link = url_for('verify_email', token=verification_token, _external=True)
    message = f'Please verify your email by clicking the following link: {verification_link}'

    msg = MIMEText(message)
    msg['Subject'] = 'Email Verification'
    msg['From'] = 'bayyinhelp@gmail.com'
    msg['To'] = email

    try:
        with smtplib.SMTP('smtp.gmail.com', 587) as server:  # Use your SMTP server
            server.starttls()
            server.login('bayyinhelp@gmail.com', 'kgrz otqt rckn pnwv')
            server.sendmail(msg['From'], [msg['To']], msg.as_string())
    except Exception as e:
        print(f"Error sending email: {e}")


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        first_name = request.form['first_name']
        second_name = request.form['second_name']
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']

        # Check if all fields are filled
        if not (first_name and second_name and username and email and password):
            return render_template('regestr.html', error_message="يرجى ملء جميع الحقول المطلوبة.")

        # Check email and password pattern validation
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            return render_template('regestr.html', error_message="يرجى إدخال بريد إلكتروني صحيح.")
        if not re.match(r"(?=.\d)(?=.[!@#$%^&])[A-Za-z\d!@#$%^&]{10,}", password):
            return render_template('regestr.html',
                                   error_message="يجب أن تحتوي كلمة المرور على 10 أحرف على الأقل، ويجب أن تتضمن رقمًا واحدًا ورمزًا خاصًا.")

        password_hashed = generate_password_hash(password)
        verification_token = str(uuid.uuid4())  # Generate verification token

        try:
            cursor = mysql.connection.cursor()
            cursor.execute(
                'INSERT INTO tbl_users (first_name, second_name, username, email, password, verification_token) VALUES (%s, %s, %s, %s, %s, %s)',
                (first_name, second_name, username, email, password_hashed, verification_token)
            )
            mysql.connection.commit()

            # Send verification email
            send_verification_email(email, verification_token)

            # Show confirmation alert and render the registration page with a success message
            return render_template('regestr.html',
                                   success_message="تم إرسال بريد إلكتروني للتحقق. يرجى النقر عليه للتأكيد .")
        except Exception as e:
            print(f"Error during registration: {e}")
            return f"Registration failed: {str(e)}"

    return render_template('regestr.html')


@app.route('/logout', methods=['POST'])
def logout():
    session.pop('loggedin', None)
    session.pop('username', None)
    session.pop('id', None)
    return redirect(url_for('login'))

def get_user_results(user_id):
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)

    try:
        # Fetch both text and dataset results
        cursor.execute("""
            SELECT 
                id,
                original_filename AS display_filename,
                'text_analysis' AS processing_type,
                '0' AS gpt4_datasets_id,
                '0' AS claude_datasets_id,
                '0' AS gemini_datasets_id,
                created_at
            FROM tbl_results
            WHERE user_id = %s AND original_text IS NOT NULL
            UNION
            SELECT 
                id,
                dataset_name AS display_filename,
                'api_analysis' AS processing_type,
                '0' AS gpt4_datasets_id,
                '0' AS claude_datasets_id,
                '0' AS gemini_datasets_id,
                created_at
            FROM tbl_datasets
            WHERE user_id = %s AND model_used IS NULL
            UNION 
            SELECT 
                id, 
                original_filename AS display_filename, 
                'ai_verifications' AS processing_type,
                '0' AS gpt4_datasets_id,
                '0' AS claude_datasets_id,
                '0' AS gemini_datasets_id,
                created_at
            FROM tbl_data_verifications
            WHERE user_id = %s AND dataset_issues IS NOT NULL
            UNION
            SELECT 
                id,
                original_text AS display_filename,
                'zero_shot' AS processing_type,
                gpt4_datasets_id,
                claude_datasets_id,
                gemini_datasets_id,
                created_at
            FROM tbl_api_analysis
            WHERE user_id = %s 
            ORDER BY created_at DESC
        """, (user_id, user_id, user_id, user_id))

        return cursor.fetchall()
    except Exception as e:
        print(f"Error in get_user_results: {e}")
        return []
    finally:
        cursor.close()



# Home page route
@app.route('/home')
def home():
    if not session.get('loggedin'):
        return render_template("homeBage.html", recent_works=None, message="يرجى تسجيل الدخول أولاً.")

    user_id = session['id']

    # جلب نتائج المستخدم
    user_results = get_user_results(user_id)

    if not user_results:
        return render_template("homeBage.html", recent_works=[], message="لا توجد أعمال بعد.")

    # ترتيب النتائج من الأحدث للأقدم وأخذ آخر 3
    recent_works = sorted(user_results, key=lambda x: x['created_at'], reverse=True)[:4]

    return render_template("homeBage.html", recent_works=recent_works)



@app.route('/terms')
def terms():
    return render_template('terms.html')


@app.route('/faq')
def faq():
    return render_template('faq.html')


from flask import send_file  # Ensure this import is included

# Upload and process file route
# Upload and process file route
from flask import render_template

import uuid
from openai import OpenAI
from flask import render_template, request, redirect, url_for, session
import os

# Initialize OpenAI client


from openai import OpenAI
from flask import render_template, request, redirect, url_for, session
import os

from openai import OpenAI

# Initialize OpenAI client (add your API key here)
client = OpenAI(
    api_key="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")


def chunk_text(text, max_chunk_size=4000):
    """Split text into smaller chunks."""
    words = text.split()
    chunks = []
    current_chunk = []
    current_size = 0

    for word in words:
        current_size += len(word) + 1  # +1 for space
        if current_size > max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_size = len(word)
        else:
            current_chunk.append(word)

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks


import os
import uuid
from flask import Flask, render_template, request, redirect, url_for, session, flash
from werkzeug.utils import secure_filename
from flask_mysqldb import MySQL

import os
import uuid
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify
from werkzeug.utils import secure_filename
from flask_mysqldb import MySQL



import os
import uuid
from flask import session, request, redirect, url_for, render_template, flash, jsonify
from werkzeug.utils import secure_filename


@app.route('/upload', methods=['POST', 'GET'])
def upload_file():
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        try:
            print("📌 Received POST request to /upload")

            selected_previous_file = request.form.get('selected_previous_file')
            print(f"🔍 Selected Previous File: {selected_previous_file}")

            # Case 1: User selects a previously uploaded file
            if selected_previous_file:
                cursor = mysql.connection.cursor()
                cursor.execute('''
                    SELECT filename, original_filename 
                    FROM tbl_results 
                    WHERE filename = %s AND user_id = %s
                ''', (selected_previous_file, session.get('id')))
                file_info = cursor.fetchone()
                cursor.close()

                if file_info:
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_info[0])
                    if os.path.exists(file_path):
                        session['uploaded_files'] = [file_path]
                        session.modified = True  # ✅ Ensure session is updated
                        print(f"✅ Previous file found and selected: {file_path}")

                        # Redirect based on process type
                        process_type = session.get('process_type', 'processing_options')
                        return redirect(url_for(process_type))

                flash('⚠️ The selected file does not exist.', 'error')
                print(f"⚠️ The selected file does not exist.")
                return render_template('upload.html')

            # Case 2: User uploads new files
            if 'files' not in request.files:
                flash('❌ No file uploaded.', 'error')
                print(f"❌ No file uploaded.")

                return render_template('upload.html')

            files = request.files.getlist('files')
            if not files or all(file.filename.strip() == '' for file in files):
                flash('❌ No file uploaded.', 'error')
                print(f"❌ No file uploaded 2.")

                return render_template('upload.html')

            uploaded_files = []
            original_filename = ''
            uuid_filename = ''

            cursor = mysql.connection.cursor()

            process_type = session.get('process_type', 'processing_options')

            # Determine allowed file extensions
            allowed_extensions = ['.xls', '.xlsx', '.csv'] if process_type in ['ai_verification', 'human_verification'] else ['.txt', '.docx']


            for file in files:
                if file and file.filename:
                    # original_filename = secure_filename(file.filename)  # Prevent path traversal
                    original_filename = file.filename  # Prevent path traversal

                    extension = os.path.splitext(original_filename)[1].lower()

                    if extension not in allowed_extensions:
                        flash('⚠️ Unsupported file format.', 'error')
                        return render_template('upload.html')

                    uuid_filename = f"{uuid.uuid4()}{extension}"
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], uuid_filename)
                    file.save(filepath)
                    uploaded_files.append(filepath)

                    cursor.execute('''
                        INSERT INTO tbl_results (user_id, filename, original_filename)
                        VALUES (%s, %s, %s)
                    ''', (session.get('id'), uuid_filename, original_filename))

            mysql.connection.commit()
            cursor.close()

            # ✅ Store file paths and original name in session before redirecting
            session['uploaded_files'] = uploaded_files
            session['unique_file_name'] = uuid_filename
            session['original_filename'] = original_filename
            session.modified = True  # ✅ Ensure session is updated before redirect

            # ✅ Redirect if AI verification is selected
            if process_type == 'ai_verification':
                print(f"🔄 Redirecting to AI verification with files: {session.get('uploaded_files')}")
                return redirect(url_for('ai_verification'))

            # ✅ Redirect to the correct processing page
            valid_routes = ['processing_options', 'api_analysis', 'human_verification', 'ai_verification']
            if process_type in valid_routes:
                return redirect(url_for(process_type))

            flash('⚠️ Invalid processing type.', 'error')
            return render_template('upload.html')

        except Exception as e:
            print(f"❌ Error during file upload: {e}")
            return jsonify({"error": f"❌ Error during file upload: {str(e)}"})

    return render_template('upload.html')


@app.route('/get_previous_files')
def get_previous_files():
    if 'loggedin' not in session:
        return jsonify({"files": [], "error": "User not logged in"})

    try:
        cursor = mysql.connection.cursor()
        # Simplified query to get just the essential information
        cursor.execute("""
            SELECT filename, original_filename 
            FROM tbl_results 
            WHERE user_id = %s 
            ORDER BY id DESC
        """, (session['id'],))

        files = cursor.fetchall()

        formatted_files = []
        seen_filenames = set()  # Track unique filenames

        for file in files:
            if file[0] and file[1]:  # Make sure both filename and original_filename exist
                uuid_filename = file[0]
                original_filename = file[1]

                # Only add if we haven't seen this original filename before
                if original_filename not in seen_filenames:
                    seen_filenames.add(original_filename)
                    formatted_files.append({
                        "filename": uuid_filename,
                        "display_name": original_filename
                    })

        cursor.close()

        # Debug print
        print("Found files:", formatted_files)

        return jsonify({"files": formatted_files})

    except Exception as e:
        print(f"Error fetching files: {e}")
        return jsonify({"files": [], "error": str(e)})


@app.route('/process_file', methods=['POST'])
def process_file():
    from flask import request, jsonify

    try:
        # Get the selected file from the request
        data = request.get_json()
        selected_file = data.get('file')

        if not selected_file:
            return jsonify({"success": False, "error": "No file selected."})

        # Process the file (e.g., read and analyze it)
        # Replace the logic below with your actual file processing
        print(f"Processing file: {selected_file}")

        # Simulate file processing
        # ... Your processing logic here ...

        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/choose_processing', methods=['GET', 'POST'])
def choose_processing():
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        process_type = request.form.get('process_type')
        session['process_type'] = process_type

        if process_type == 'zero_shot':
            return redirect(url_for('zero_shot'))  # Or any correct endpoint

        else:
            return redirect(url_for('upload_file'))
    return render_template('choose_processing.html')


@app.route('/zero_shot', methods=['GET'])
def zero_shot():
    if 'loggedin' not in session:
        return redirect(url_for('login'))
    return render_template('zero_shot.html')


from scipy import stats



from scipy import stats


def detect_unique_values(df):
    """يرجع عدد القيم الفريدة في كل عمود."""
    return df.nunique().to_dict()


def detect_missing_values(df):
    """يرجع الأعمدة التي تحتوي على قيم مفقودة وعددها."""
    missing = df.isnull().sum()
    return {col: int(count) for col, count in missing.items() if count > 0} or "✅ لا توجد قيم مفقودة."


def detect_duplicate_rows(df):
    """يرجع عدد الصفوف المكررة إن وجدت."""
    duplicate_count = df.duplicated().sum()
    return {"عدد الصفوف المكررة": int(duplicate_count)} if duplicate_count > 0 else "✅ لا توجد صفوف مكررة."


from scipy.stats import zscore

def detect_outliers(df):
    """يرجع الأعمدة الرقمية التي تحتوي على قيم متطرفة."""
    outlier_columns = []

    for col in df.select_dtypes(include=['number']).columns:
        if df[col].nunique() <= 1:
            continue

        z_scores = zscore(df[col], nan_policy='omit')
        if (abs(z_scores) > 3).any():
            outlier_columns.append(col)

    return outlier_columns if outlier_columns else "✅ لا توجد قيم متطرفة."


def check_column_data_types(df):
    """يرجع الأعمدة التي تحتوي على قيم غير متوافقة مع نوع البيانات المتوقع."""
    type_issues = {}

    for col in df.columns:
        if df[col].dtype == 'object' and df[col].str.isnumeric().sum() > 0:
            type_issues[col] = "⚠ يحتوي على أرقام رغم أنه عمود نصي."
        elif df[col].dtype in ['int64', 'float64'] and df[col].apply(lambda x: isinstance(x, str)).sum() > 0:
            type_issues[col] = "⚠ يحتوي على نصوص رغم أنه عمود رقمي."

    return type_issues if type_issues else "✅ جميع الأعمدة تحتوي على أنواع بيانات صحيحة."


def detect_category_imbalance(df, threshold=0.90):
    """يرجع أسماء الأعمدة الفئوية التي فيها فئة واحدة تهيمن على أكثر من 90٪ من القيم."""
    imbalanced_columns = []

    for col in df.select_dtypes(include=['object']).columns:
        if df[col].nunique() <= 1:
            continue
        top_category_ratio = df[col].value_counts(normalize=True).values[0]
        if top_category_ratio > threshold:
            imbalanced_columns.append(col)

    return imbalanced_columns if imbalanced_columns else "✅ لا يوجد عدم توازن في الأعمدة الفئوية."


import pandas as pd
import numpy as np
from scipy.stats import zscore

def validate_dataset(file_path):
    try:
        df = pd.read_csv(file_path) if file_path.endswith('.csv') else pd.read_excel(file_path)

        results = {}

        # 1. القيم الفريدة
        results["عدد القيم الفريدة"] = df.nunique().to_dict()

        # 2. القيم المفقودة
        missing = df.isnull().sum()
        results["القيم المفقودة"] = {col: int(val) for col, val in missing.items() if val > 0} or "✅ لا توجد قيم مفقودة."

        # 3. الصفوف المكررة
        results["الصفوف المكررة"] = {"عدد الصفوف المكررة": int(df.duplicated().sum())} if df.duplicated().sum() > 0 else "✅ لا توجد صفوف مكررة."

        # 4. القيم المتطرفة
        outliers = []
        for col in df.select_dtypes(include=[np.number]):
            if df[col].nunique() <= 1:
                continue
            z_scores = zscore(df[col], nan_policy='omit')
            if np.any(np.abs(z_scores) > 3):
                outliers.append(col)
        results["القيم المتطرفة"] = outliers or "✅ لا توجد قيم متطرفة."

        # 5. مشاكل نوع البيانات
        type_issues = {}
        for col in df.columns:
            if df[col].dtype == 'object' and df[col].astype(str).str.isnumeric().sum() > 0:
                type_issues[col] = "⚠ يحتوي على أرقام رغم أنه عمود نصي."
            elif df[col].dtype in ['int64', 'float64'] and df[col].apply(lambda x: isinstance(x, str)).sum() > 0:
                type_issues[col] = "⚠ يحتوي على نصوص رغم أنه عمود رقمي."
        results["مشاكل نوع البيانات"] = type_issues or "✅ لا توجد مشاكل في نوع البيانات."

        # 6. اختلال توازن الفئات
        imbalance = []
        for col in df.select_dtypes(include=['object']).columns:
            top_ratio = df[col].value_counts(normalize=True, dropna=False).max()
            if top_ratio > 0.9:
                imbalance.append(f"{col}: {top_ratio:.2%}")
        results["اختلال الفئات"] = imbalance or "✅ لا يوجد اختلال."

        # 7. مشاكل التواريخ
        date_issues = {}
        for col in df.columns:
            try:
                parsed = pd.to_datetime(df[col], errors='coerce', dayfirst=True)
                invalid = parsed.isna().sum()
                future = (parsed > pd.Timestamp.today()).sum()
                old = (parsed < '1900-01-01').sum()
                if invalid > 0 or future > 0 or old > 0:
                    date_issues[col] = {
                        "تواريخ غير صالحة": int(invalid),
                        "تواريخ مستقبلية": int(future),
                        "تواريخ قديمة جدًا": int(old)
                    }
            except:
                continue
        results["مشاكل التواريخ"] = date_issues or "✅ لا توجد مشاكل في التواريخ."

        return results

    except Exception as e:
        return {"خطأ": str(e)}


def validate_file(file_path):
    """يشغل كل اختبارات التحقق ويجمع النتائج."""
    try:
        df = pd.read_csv(file_path) if file_path.endswith('.csv') else pd.read_excel(file_path, engine='openpyxl')
        if df.empty:
            return {"خطأ": "❌ الملف فارغ أو تالف."}

        detected_issues = {}

        detected_issues["عدد القيم الفريدة"] = detect_unique_values(df)
        detected_issues["القيم المفقودة"] = detect_missing_values(df)
        detected_issues["الصفوف المكررة"] = detect_duplicate_rows(df)
        detected_issues["القيم المتطرفة"] = detect_outliers(df)
        detected_issues["مشاكل في نوع البيانات"] = check_column_data_types(df)
        detected_issues["عدم توازن الفئات"] = detect_category_imbalance(df)

        return detected_issues

    except Exception as e:
        return {"خطأ": f"❌ تعذر قراءة الملف: {str(e)}"}


import os
import json
import pandas as pd
from flask import Flask, request, jsonify, render_template
from openai import OpenAI

# ✅ Function to Convert Non-Serializable Data to JSON-Safe Format
def convert_to_json_safe(data):
    """Recursively converts non-serializable types (e.g., int64) to Python-native types."""
    if isinstance(data, dict):
        return {k: convert_to_json_safe(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [convert_to_json_safe(v) for v in data]
    elif isinstance(data, (np.int64, np.float64, int, float)):
        return int(data) if isinstance(data, (np.int64, int)) else float(data)
    elif isinstance(data, pd.Timestamp):
        return str(data)  # Convert Timestamp to string
    else:
        return data


# ✅ Function to Detect Issues in Dataset
def detect_issues(df):
    """Detects missing values, duplicates, outliers, data type issues, and category imbalances."""
    issues = {}

    # 🔴 Detect Missing Values
    missing_values = df.isnull().sum().to_dict()
    missing_values = {col: int(count) for col, count in missing_values.items() if count > 0}
    issues["Missing Values"] = missing_values if missing_values else {}

    # 🔴 Detect Duplicate Rows
    duplicate_count = int(df.duplicated().sum())
    issues["Duplicates Found"] = {"Total Duplicates": duplicate_count} if duplicate_count > 0 else {}

    # 🔴 Detect Outliers (Z-score method)
    outliers = {}
    for col in df.select_dtypes(include=['number']).columns:
        if df[col].nunique() > 1:
            z_scores = stats.zscore(df[col], nan_policy='omit')
            outlier_indices = df.index[abs(z_scores) > 3].tolist()
            if outlier_indices:
                outliers[col] = [int(idx) for idx in outlier_indices]
    issues["Outliers Detected"] = outliers if outliers else {}

    # 🔴 Detect Data Type Issues
    data_type_issues = {}
    for col in df.columns:
        if df[col].dtype == 'object' and df[col].str.isnumeric().sum() > 0:
            data_type_issues[col] = "⚠ Unexpected numeric values found."
        elif df[col].dtype in ['int64', 'float64'] and df[col].apply(lambda x: isinstance(x, str)).sum() > 0:
            data_type_issues[col] = "⚠ Unexpected text values found."
    issues["Data Type Issues"] = data_type_issues if data_type_issues else {}

    # 🔴 Detect Date Validation Issues
    date_issues = {}
    for col in df.select_dtypes(include=['object']).columns:
        try:
            df[col] = pd.to_datetime(df[col], errors='coerce', dayfirst=True)
            invalid_dates = df[df[col].isna()].index.tolist()
            future_dates = df[df[col] > pd.Timestamp.today()].index.tolist()
            impossible_dates = [i for i, date in df[col].dropna().items() if (date.day == 31 and date.month == 2) or (date.day > 29 and date.month == 2)]

            date_issues[col] = {
                "Invalid Dates": invalid_dates if invalid_dates else "✅ No invalid dates detected.",
                "Impossible Dates": impossible_dates if impossible_dates else "✅ No impossible dates detected.",
                "Future Dates": future_dates if future_dates else "✅ No future dates detected."
            }
        except Exception as e:
            date_issues[col] = {"error": f"Error in date validation: {str(e)}"}

    issues["Date Validation Issues"] = date_issues if date_issues else {}

    # 🔴 Detect Category Imbalances
    imbalance_report = {}
    for col in df.select_dtypes(include=['object']).columns:
        top_category_ratio = df[col].value_counts(normalize=True).max()
        if top_category_ratio > 0.90:
            imbalance_report[col] = f"⚠ Dominant category covers {top_category_ratio:.2%} of data."
    issues["Category Imbalance"] = imbalance_report if imbalance_report else {}

    return issues






# ✅ Function to Get GPT-4o Recommendations
def get_gpt_recommendations(issues):
    """استخدام GPT-4o لإعطاء توصيات باللغة العربية بناءً على المشاكل المكتشفة في البيانات."""

    # 🔹 إزالة الأقسام التي لا تحتوي على مشاكل
    filtered_issues = {key: value for key, value in issues.items() if value and value != "✅ No issues detected."}

    if not filtered_issues:
        return "✅ لا توجد مشاكل تتطلب توصيات."

    issues_json = json.dumps(filtered_issues, indent=2, ensure_ascii=False)
    max_chars = 50000
    issues_limited = issues_json[:max_chars]

    # ✅ Arabic Prompt
    prompt = f"""
    هذه نتائج تحليل جودة البيانات:

    {issues_limited}

    🔧 أرجو أن تقدم توصيات واضحة لتحسين جودة البيانات، وذلك وفقاً للنقاط التالية:

    1. ✨ التوصيات يجب أن تكون *مكتوبة باللغة العربية الفصحى* وبأسلوب احترافي ومباشر.
    2. ✨ كل مشكلة يجب أن تُتبع بتوصية واحدة أو أكثر لحلّها، بدون تكرار أو تعميم.
    3. ✨ استخدم علامات تنسيق مثل (نقطة - عنوان فرعي) عند الحاجة لتقسيم الأفكار.
    4. ✨ لا تُدرج أي معلومات عامة أو غير موجودة في المشاكل المذكورة.
    5. ✨ لا تذكر "المشكلة" مرة أخرى، فقط ابدأ بالتوصية كأنك تكتب تقريراً عملياً للمستخدم.

    🎯 الهدف هو إعطاء المستخدم خطوات قابلة للتنفيذ لتحسين جودة ملف البيانات.

    اكتب التوصيات الآن:
    """

    try:
        response = gpt_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=600
        )
        return response.choices[0].message.content.strip()

    except Exception as e:
        return f"⚠ خطأ أثناء جلب التوصيات: {str(e)}"
from scipy.stats import zscore

# ✅ Flask Route for view verification result
@app.route('/view_ai_verification/<int:dataset_id>')
def view_ai_verification(dataset_id):
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute("""
            SELECT *
            FROM tbl_data_verifications
            WHERE id = %s AND user_id = %s
        """, (dataset_id, session['id']))

        result = cursor.fetchone()

        if not result:
            return render_template('view_dataset.html', error="البيانات غير موجودة.")

        # ✅ Ensure JSON Parsing is handled safely
        try:
            dataset_issues = json.loads(result['dataset_issues'])
            recommendations = result['recommendations']
        except json.JSONDecodeError as e:
            print(f"❌ Error parsing JSON: {e}")
            return render_template('view_dataset.html', error="خطأ في قراءة البيانات.")

        # ✅ Render dataset details in the template
        return render_template('dataset_quality_page.html', results=dataset_issues, recommendations=recommendations)


    except Exception as e:
        print(f"❌ Error viewing dataset: {e}")
        return render_template('view_dataset.html', error="حدث خطأ أثناء تحميل البيانات.")

    finally:
        cursor.close()


# ✅ Flask Route for Dataset Validation
# ✅ Flask Route for Dataset Validation
@app.route('/ai_verification', methods=['GET', 'POST'])
def ai_verification():
    uploaded_files = session.get('uploaded_files', [])

    if not uploaded_files:
        print("🚨 No file found in session!")
        return jsonify({'error': '❌ No file found in session.'}), 400

    file_path = uploaded_files[0]  # ✅ Retrieve the first uploaded file

    try:
        df = pd.read_csv(file_path) if file_path.endswith('.csv') else pd.read_excel(file_path, engine='openpyxl')
        if df.empty:
            return jsonify({'error': '❌ The uploaded dataset is empty or corrupted.'}), 400

        dataset_issues = {
            "Unique Values Count per Column": detect_unique_values(df),
            "Missing Values": detect_missing_values(df),
            "Duplicate Rows": detect_duplicate_rows(df),
            "Outliers Detected": detect_outliers(df),
            "Data Type Issues": check_column_data_types(df),
            "Category Imbalance": detect_category_imbalance(df)
        }

        print("🔍 DEBUGGING OUTPUT:", json.dumps(dataset_issues, indent=2, ensure_ascii=False))
        recommendations = get_gpt_recommendations(dataset_issues)

        cursor = mysql.connection.cursor()

        cursor.execute('''
                    INSERT INTO tbl_data_verifications (user_id, dataset_issues, recommendations, type, filename, original_filename)
                    VALUES (%s, %s, %s, %s, %s, %s)
                ''', (session.get('id'), json.dumps(dataset_issues, ensure_ascii=False), recommendations, 'ai_verification', session.get('unique_file_name'), session.get('original_filename')))

        mysql.connection.commit()
        cursor.close()

        return render_template('dataset_quality_page.html', results=dataset_issues, recommendations=recommendations)

    except Exception as e:
        print(f"❌ Error processing file: {e}")
        return jsonify({'error': f'❌ Error processing file: {str(e)}'}),500



def process_text_file(text):
    """Processes large text in smaller chunks to avoid GPT-4o response limits."""
    chunk_size = 4000  # Fits safely within GPT-4o context window
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    full_data = {
        "نوع_المحتوى": "غير محدد",
        "الملخص": "",
        "الأعمدة": [],
        "البيانات": []
    }

    for chunk in chunks:
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": """
📄 تحليل نص غير منظم وتحويله إلى بيانات منظمة بصيغة JSON.

⚙️ لا يوجد مجال محدد للنص (قد يكون تقريرًا، مقالًا، مقابلة، بيانات خام... إلخ).

🎯 المهام المطلوبة:
1.⁠ ⁠استنتاج نوع المحتوى بدقة (مثال: تقرير مقابلات، تحليل مالي، مقال ثقافي...).
2.⁠ ⁠تقديم ملخص قصير وواضح يشرح فكرة النص أو هدفه.
3.⁠ ⁠استنتاج الأعمدة (Columns) المنطقية المناسبة من النص بدون فرض عدد محدد.
   - استخدم أسماء أعمدة وصفية وتعكس معنى المعلومات.
   - تجنب التكرار في أسماء الأعمدة.
4.⁠ ⁠إنشاء صفوف نموذجية (Rows) تعكس البنية الصحيحة للبيانات (كل صف يحتوي على نفس عدد الأعمدة).
   - لا تضف صفوف ناقصة أو مكررة.
   - تجاهل أي محتوى غير قابل للتنظيم في جدول.

📦 الإخراج بصيغة JSON فقط، كالتالي:
{
  "نوع_المحتوى": "...",
  "الملخص": "...",
  "الأعمدة": ["...", "...", "..."],
  "البيانات": [["...", "...", "..."], ["...", "...", "..."]]
}
"""
                    },
                    {
                        "role": "user",
                        "content": f"🔹 تحليل هذا الجزء من النص:\n{chunk}"
                    }
                ],
                response_format={"type": "json_object"},
                temperature=0.2,
                max_tokens=4096
            )

            result_text = response.choices[0].message.content
            result_dict = json.loads(result_text)

            # Set or update common metadata
            full_data["نوع_المحتوى"] = result_dict.get("نوع_المحتوى", full_data["نوع_المحتوى"])
            if result_dict.get("الملخص"):
                full_data["الملخص"] = result_dict["الملخص"].strip()

            # Set columns if empty (from first chunk only)
            if not full_data["الأعمدة"] and "الأعمدة" in result_dict:
                full_data["الأعمدة"] = result_dict["الأعمدة"]

            # Append rows if valid
            if "البيانات" in result_dict:
                for row in result_dict["البيانات"]:
                    if isinstance(row, list) and len(row) == len(full_data["الأعمدة"]):
                        full_data["البيانات"].append(row)

        except Exception as e:
            print(f"❌ Error processing chunk: {e}")

    return full_data


# ✅ Now define api_analysis() AFTER process_text_file()
@app.route('/api_analysis', methods=['GET'])
def api_analysis():
    """Processes the uploaded file and sends it to GPT-4o for dataset generation."""
    if 'loggedin' not in session or 'uploaded_files' not in session:
        return redirect(url_for('login'))

    try:
        merged_text = ""

        # ✅ Read uploaded file
        for file_path in session.get('uploaded_files', []):
            print(f"🧾 Trying to read: {file_path} | Exists? {os.path.exists(file_path)}")

            content = ""  # تأكد أن المتغير معرف

            try:
                # أول محاولة بالترميز UTF-8
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read().strip()
            except UnicodeDecodeError:
                print(f"⚠️ UTF-8 failed for {file_path}, trying windows-1256...")
                try:
                    with open(file_path, 'r', encoding='windows-1256') as file:
                        content = file.read().strip()
                except Exception as nested_error:
                    print(f"❌ Encoding fallback failed: {nested_error}")
                    return render_template('api_results.html', error="❌ تعذر قراءة الملف بسبب مشكلة في الترميز.")

            except Exception as file_error:
                print(f"❌ File Read Error: {file_error}")
                return render_template('api_results.html', error="❌ حدث خطأ أثناء قراءة الملف.")

            if content:
                print(f"📄 First 100 characters: {content[:100]}")
                merged_text += content + "\n\n"
            else:
                print(f"⚠️ Empty or unreadable content in: {file_path}")

        # ✅ Ensure text exists
        if not merged_text.strip():
            return render_template('api_results.html', error="❌ لم يتم العثور على نص داخل الملف.")

        # ✅ Send to GPT-4o in chunks
        result = process_text_file(merged_text)

        print("✅ Final Processed Data:", json.dumps(result, ensure_ascii=False, indent=4))

        # ✅ Ensure valid extracted data
        if not result or not result.get('البيانات'):
            return render_template('api_results.html', error="❌ لم يتم استخراج بيانات منظمة من النص.")

        extracted_columns = result.get('الأعمدة', [])
        extracted_rows = result.get('البيانات', [])
        extracted_summary = result.get('الملخص', "لا يوجد ملخص متاح")
        extracted_content_type = result.get('نوع_المحتوى', "غير محدد")

        # ✅ Store in database
        cursor = mysql.connection.cursor()
        cursor.execute("""
            INSERT INTO tbl_datasets 
            (user_id, dataset_name, content_type, columns_data, rows_data, original_text, summary) 
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (
            session['id'],
            "مجموعة بيانات منظمة",
            extracted_content_type,
            json.dumps(extracted_columns, ensure_ascii=False),
            json.dumps(extracted_rows, ensure_ascii=False),
            merged_text,
            extracted_summary
        ))

        dataset_id = cursor.lastrowid
        mysql.connection.commit()

        return render_template('api_results.html',
                               results=[result],
                               result_id=dataset_id,
                               original_text=merged_text,
                               column_count=len(extracted_columns),
                               row_count=len(extracted_rows),
                               summary=extracted_summary,
                               content_type=extracted_content_type)

    except Exception as e:
        print(f"❌ Error in API analysis: {e}")
        return render_template('api_results.html', error=f"❌ حدث خطأ أثناء التحليل: {str(e)}")


def extract_basic_info(text):
    """Extract basic information from text when structured parsing fails."""
    # Split the text into parts and take the first few elements
    parts = text.split('،')
    result = []

    # Try to extract at least person and position
    if len(parts) >= 2:
        result = [parts[0].strip(), parts[1].strip(), ' '.join(parts[2:]).strip()]
    else:
        result = [text.strip(), "", ""]

    return result




@app.route('/view_dataset/<int:dataset_id>')
def view_dataset(dataset_id):
    """Fetch dataset details from the database and render them properly."""
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute("""
            SELECT id, dataset_name, content_type, columns_data, rows_data, 
                   created_at, original_text, summary 
            FROM tbl_datasets 
            WHERE id = %s AND user_id = %s
        """, (dataset_id, session['id']))

        result = cursor.fetchone()

        if not result:
            return render_template('view_dataset.html', error="البيانات غير موجودة.")

        # ✅ Ensure JSON Parsing is handled safely
        try:
            columns = json.loads(result['columns_data'])
            rows = json.loads(result['rows_data'])
        except json.JSONDecodeError as e:
            print(f"❌ Error parsing JSON: {e}")
            return render_template('view_dataset.html', error="خطأ في قراءة البيانات.")

        # ✅ Render dataset details in the template
        return render_template('view_dataset.html',
                               dataset={
                                   'id': result['id'],
                                   'name': result['dataset_name'],
                                   'content_type': result['content_type'],
                                   'columns': columns,
                                   'rows': rows,
                                   'created_at': result['created_at'],
                                   'original_text': result['original_text'],
                                   'summary': result['summary']
                               })

    except Exception as e:
        print(f"❌ Error viewing dataset: {e}")
        return render_template('view_dataset.html', error="حدث خطأ أثناء تحميل البيانات.")

    finally:
        cursor.close()


app.config["UPLOAD_FOLDER"] = "generated_datasets"  # Ensure this folder exists

if not os.path.exists(app.config["UPLOAD_FOLDER"]):
    os.makedirs(app.config["UPLOAD_FOLDER"])


import os
import json
import MySQLdb
import anthropic  # ✅ Importing Anthropic for Claude 3
from flask import Flask, render_template, request, jsonify, session, url_for
from openai import OpenAI  # ✅ Make sure you have `openai` installed
from dotenv import load_dotenv

load_dotenv()  # ✅ Load API keys from `.env` (recommended)



# ✅ Load API Keys from Environment Variables (Recommended)
@app.route('/export_dataset/<int:dataset_id>/<format_type>')
def export_dataset(dataset_id, format_type):
    """Exports dataset as CSV or XLSX."""
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute("""
            SELECT columns_data, rows_data, dataset_name 
            FROM tbl_datasets 
            WHERE id = %s AND user_id = %s
        """, (dataset_id, session['id']))
        result = cursor.fetchone()

        if not result:
            return "Dataset not found", 404

        columns = json.loads(result['columns_data'])
        rows = json.loads(result['rows_data'])

        if format_type.lower() == 'csv':
            output = StringIO()
            writer = csv.writer(output)

            # Write headers
            writer.writerow(columns)

            # Write data rows
            writer.writerows(rows)

            return Response(
                output.getvalue().encode('utf-8-sig'),
                mimetype='text/csv',
                headers={"Content-Disposition": f"attachment;filename=dataset_{dataset_id}.csv"}
            )

        elif format_type.lower() == 'xlsx':
            output = BytesIO()
            workbook = xlsxwriter.Workbook(output)
            worksheet = workbook.add_worksheet()

            # Write headers
            for col, header in enumerate(columns):
                worksheet.write(0, col, header)

            # Write data rows
            for row_idx, row_data in enumerate(rows, 1):
                for col_idx, cell_value in enumerate(row_data):
                    worksheet.write(row_idx, col_idx, cell_value)

            workbook.close()
            output.seek(0)

            return Response(
                output.getvalue(),
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                headers={"Content-Disposition": f"attachment;filename=dataset_{dataset_id}.xlsx"}
            )

    except Exception as e:
        print(f"Error exporting dataset: {e}")
        return "Error exporting dataset", 500

import mimetypes
import os
import json
import MySQLdb
from flask import Flask, render_template, request, jsonify, session, url_for
from openai import OpenAI
import anthropic
@app.route('/human_verification', methods=['GET'])
def human_verification():
    # Your verification logic
    return jsonify({"success": True, "message": "Human verification processed successfully."})
import pandas as pd
import uuid

UPLOAD_FOLDER = "static/uploads"  # ✅ Store files in a public directory
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)  # ✅ Ensure the upload directory exists


from urllib.parse import quote

@app.route('/send_email', methods=['POST'])
def send_email():
    try:
        email = request.form.get("email")
        message = request.form.get("message")
        file = request.files.get("file")

        if not email or not message or not file:
            return jsonify({"success": False, "error": "❌ البريد الإلكتروني أو الرسالة أو الملف مفقود!"})

        # إعدادات البريد
        sender_email = "bayyinhelp@gmail.com"
        sender_password = "kgrz otqt rckn pnwv"  # استخدم App Password

        # حفظ الملف
        file_id = str(uuid.uuid4())
        filename = f"{file_id}_{file.filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # ✅ Encode only for the URL (not for database)
        encoded_message = quote(message)

        # إنشاء رابط المعاينة مع تمرير الرسالة كمعلومة في الرابط
        dataset_link = url_for('preview_dataset', file_id=file_id, filename=file.filename, _external=True) + f"?message={encoded_message}"

        # نص الرسالة في الإيميل
        email_body = f"""
        <html>
        <body style="font-family: Arial, sans-serif;">
            <p style="white-space: pre-wrap;">{message}</p>
            <p>📂 لمراجعة البيانات، اضغط على الرابط التالي:</p>
            <a href="{dataset_link}" target="_blank" style="color: #7d2ae8; font-weight: bold;">🔗 عرض البيانات</a>
        </body>
        </html>
        """

        # تجهيز الإيميل
        msg = MIMEMultipart()
        msg["From"] = sender_email
        msg["To"] = email
        msg["Subject"] = "📊 عرض البيانات المرسلة"
        msg.attach(MIMEText(email_body, "html"))

        # إرسال الإيميل
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, email, msg.as_string())

        # ✅ Save raw clean message in the database
        cursor = mysql.connection.cursor()
        cursor.execute('''
            INSERT INTO tbl_data_verifications (user_id, type, filename, original_filename, verification_email, file_id, email_message)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        ''', (
            session.get('id'),
            'human_verification',
            filename,
            file.filename,
            email,
            file_id,
            message  # ✅ Save the normal message
        ))

        mysql.connection.commit()
        cursor.close()

        return jsonify({"success": True, "message": "✅ تم إرسال البريد الإلكتروني بنجاح!"})

    except Exception as e:
        print("Error while sending email:", e)
        return jsonify({"success": False, "error": f"❌ خطأ أثناء إرسال البريد الإلكتروني: {str(e)}"})


from flask import send_file


@app.route('/preview_human_verification_rating/<int:verification_rating_id>', methods=['GET'])
def preview_human_verification_rating(verification_rating_id):
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)

        # Join both tables based on data_verification_id
        cursor.execute("""
            SELECT 
                r.*, 
                d.verification_email, 
                d.email_message 
            FROM tbl_human_verifications_rating r
            JOIN tbl_data_verifications d 
            ON r.data_verification_id = d.id
            WHERE d.user_id = %s AND r.data_verification_id = %s
        """, (session.get('id'), verification_rating_id))

        verification_rating = cursor.fetchone()

        if not verification_rating:
            return render_template('preview_human_verification_rating.html', error="❌ لم يتم ارسال التقييم بعد!")

        return render_template('preview_human_verification_rating.html', verification_rating=verification_rating)

    except Exception as e:
        print("Error previewing human verification:", e)
        return render_template('preview_human_verification_rating.html', error="❌ حدث خطأ أثناء عرض التقييم.")


@app.route('/set_human_verification_rating', methods=['POST'])
def set_human_verification_rating():
    try:
        is_data_organized = request.form.get("is_data_organized")
        is_content_clear_easy = request.form.get("is_content_clear_easy")
        evaluate_column_formatting = request.form.get("evaluate_column_formatting")
        is_data_amount_sufficient_accurate_results = request.form.get("is_data_amount_sufficient_accurate_results")
        most_liked_about_data = request.form.get("most_liked_about_data")
        comments_or_suggestions = request.form.get("comments_or_suggestions")
        overall_dataset_opinion = request.form.get("overall_dataset_opinion")
        filename = request.form.get("filename")

        user_id = session.get('id')  # 👈 إضافة user_id

        cursor = mysql.connection.cursor()

        cursor.execute("""
            SELECT id FROM tbl_data_verifications
            WHERE filename = %s
        """, (filename,))
        result = cursor.fetchone()

        if (not is_data_organized or not is_content_clear_easy or not
            evaluate_column_formatting or not is_data_amount_sufficient_accurate_results or not
            most_liked_about_data or not comments_or_suggestions or not overall_dataset_opinion):
            return jsonify({"success": False, "error": "❌ يرجى تعبئة جميع الأسئلة قبل إرسال التقييم!"})

        if (not result):
            return jsonify({"success": False, "error": "❌ الملف غير موجود!"})

        print("result['id'] ===> ", result)

        cursor.execute('''
            INSERT INTO tbl_human_verifications_rating (
                data_verification_id, 
                is_data_organized, 
                is_content_clear_easy, 
                evaluate_column_formatting, 
                is_data_amount_sufficient_accurate_results, 
                most_liked_about_data, 
                comments_or_suggestions,
                overall_dataset_opinion)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        ''', (result[0],
              is_data_organized,
              is_content_clear_easy,
              evaluate_column_formatting,
              is_data_amount_sufficient_accurate_results,
              most_liked_about_data,
              comments_or_suggestions,
              overall_dataset_opinion
        ))

        mysql.connection.commit()
        cursor.close()

        return jsonify({"success": True, "message": "✅ تم إرسال التقييم بنجاح!"})

    except Exception as e:
        print(traceback.format_exc())
        return jsonify({"success": False, "error": f"❌ خطأ أثناء إرسال التقييم: {e}"})

@app.route('/set_human_verification', methods=['POST'])
def set_human_verification():
    try:
        dataset_issues = request.form.get("dataset_issues")
        recommendations = request.form.get("recommendations")
        filename = request.form.get("filename")

        if not dataset_issues or not recommendations:
            return jsonify({"success": False, "error": "❌ تحليل البيانات او التوصيات لتحسين البيانات مفقود!"})

        cursor = mysql.connection.cursor()
        cursor.execute('''
                    UPDATE tbl_data_verifications SET dataset_issues = %s, recommendations = %s WHERE filename = %s
                ''', (dataset_issues, recommendations, filename))

        mysql.connection.commit()
        cursor.close()

        return jsonify({"success": True, "message": "✅ تم إرسال البيانات بنجاح!"})

    except Exception as e:
        return jsonify({"success": False, "error": f"❌ خطأ أثناء إرسال البريد الإلكتروني: {e}"})



@app.route('/preview_dataset/<file_id>/<filename>', methods=['GET', 'POST'])
def preview_dataset(file_id, filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_{filename}")

        if not os.path.exists(file_path):
            return "<h3>❌ الملف غير موجود! تأكد من أنك رفعت الملف بنجاح.</h3>"

        if filename.endswith(".csv"):
            df = pd.read_csv(file_path)
        elif filename.endswith(".xlsx"):
            df = pd.read_excel(file_path)
        else:
            return "<h3>❌ الملف ليس بصيغة CSV أو Excel.</h3>"

        df = df.head(100)
        data_json = df.to_json(orient="records")
        download_url = url_for('download_dataset', file_id=file_id, filename=filename, _external=True)

        # ✅ اقرأ الرسالة من الرابط ومررها للقالب
        message = request.args.get("message", "")

        return render_template("preview_dataset.html", data_json=data_json, download_url=download_url, message=message, filename=f"{file_id}_{filename}")

    except Exception as e:
        return f"<h3>❌ حدث خطأ أثناء تحميل البيانات: {e}</h3>"


# ✅ Route to Download the Full Dataset
@app.route('/download_dataset/<file_id>/<filename>')
def download_dataset(file_id, filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{file_id}_{filename}")

        # ✅ Check if file exists before sending
        if not os.path.exists(file_path):
            return "<h3>❌ الملف غير موجود! تأكد من أنك رفعت الملف بنجاح.</h3>"

        return send_file(file_path, as_attachment=True)

    except Exception as e:
        return f"<h3>❌ خطأ أثناء تنزيل الملف: {e}</h3>"


import smtplib
import os
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders


def send_email_with_attachment(receiver_email, subject, body, file_path):
    sender_email = "bayyinhelp@gmail.com"  # Your Gmail
    sender_password = "kgrz otqt rckn pnwv"  # Use your App Password

    # ✅ Create email message
    msg = MIMEMultipart()
    msg["From"] = sender_email
    msg["To"] = receiver_email
    msg["Subject"] = subject

    # ✅ Add email body
    msg.attach(MIMEText(body, "plain"))

    # ✅ Attach the file if it exists
    if file_path and os.path.exists(file_path):
        with open(file_path, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(file_path)}")
            msg.attach(part)

    try:
        # ✅ Connect to SMTP server and send email
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()  # Secure connection
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, receiver_email, msg.as_string())

        print("✅ Email sent successfully!")
        return True
    except Exception as e:
        print(f"❌ Error sending email: {e}")
        return False


# ✅ OpenAI API Client (GPT-4o)
OPENAI_API_KEY = "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"
gpt_client = OpenAI(api_key=OPENAI_API_KEY)

# ✅ Claude 3 API Client
CLAUDE_API_KEY = "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"  # Replace with actual Claude API key
claude_client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

import google.generativeai as genai

# ✅ Replace with your actual Gemini API key
genai.configure(api_key="AIzaSyD9NsiaXq73V7wdtlNMl5P8DW9sN6yQ9nc")


response = genai.GenerativeModel("gemini-1.5-pro-latest").generate_content("Hello, how can I help you?")
print(response.text)


app.config["SESSION_TYPE"] = "filesystem"  # 🔹 Store sessions in the filesystem
app.config["SESSION_PERMANENT"] = True  # 🔹 Make session persistent

# ✅ Function: Extract JSON from AI response safely
def extract_json(response_text):
    """Extracts JSON from AI response, ensuring proper formatting."""
    try:
        json_match = re.search(r"\{.*\}", response_text, re.DOTALL)  # Extract JSON part
        if not json_match:
            print("❌ Error: No valid JSON found in AI response.")
            return {"columns": [], "rows": []}

        return json.loads(json_match.group().strip())  # Parse JSON safely

    except json.JSONDecodeError as e:
        print(f"❌ JSON Decode Error: {e}")
        return {"columns": [], "rows": []}


# ✅ Function: Generate dataset using GPT-4o with pagination (Safe JSON)
def process_text_chunk(description, domain="عام", columns=5, rows=350, goal="غير محدد"):
    """Generates dataset in chunks from GPT-4o, ensuring full completion."""
    all_rows = []
    dataset = {"columns": [], "rows": []}

    for start_row in range(0, rows, 50):  # 🔹 Use 50-row batches for stability
        chunk_size = min(50, rows - start_row)
        print(f"🔹 Requesting GPT-4o: {chunk_size} rows (from {start_row})")

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": f"""قم بإنشاء مجموعة بيانات باللغة العربية وفقًا لهذه التفاصيل:
                - *الوصف*: {description}
                - *المجال*: {domain}
                - *الهدف من الإنشاء*: {goal} ✅
                - *عدد الأعمدة*: {columns}
                - *عدد الصفوف*: {chunk_size}

                📌 *الإخراج يجب أن يكون JSON فقط*:
                ```json
                {{
                    "columns": ["اسم العمود 1", "اسم العمود 2", ...],
                    "rows": [
                        ["بيانات 1", "بيانات 2", ...],
                        ["بيانات 1", "بيانات 2", ...]
                    ]
                }}
                ```"""
            }],
            response_format={"type": "json_object"},
            temperature=0.2,
            max_tokens=2048  # 🔹 Prevent response truncation
        )

        raw_output = response.choices[0].message.content
        dataset_chunk = extract_json(raw_output)

        if dataset_chunk.get("rows"):
            all_rows.extend(dataset_chunk["rows"])  # Append new rows
            dataset["columns"] = dataset_chunk["columns"]

    return {"columns": dataset["columns"], "rows": all_rows}


# ✅ Function: Generate dataset using Claude 3 with pagination (Batch Size: 20)
def generate_claude_dataset(description, domain, columns, rows, goal="غير محدد"):
    """Generates dataset in chunks from Claude 3 while ensuring consistent columns."""
    all_rows = []
    dataset = {"columns": [], "rows": []}
    BATCH_SIZE = 20  # ✅ Keep batch size small to avoid failure
    MAX_RETRIES = 2  # ✅ Reduce retries for stability
    first_batch = True  # Track the first successful batch

    for start_row in range(0, rows, BATCH_SIZE):
        chunk_size = min(BATCH_SIZE, rows - start_row)
        print(f"🔹 Requesting Claude 3: {chunk_size} rows (from {start_row})")

        prompt = f"""
        أنت نموذج ذكاء اصطناعي. قم بإنشاء مجموعة بيانات باللغة العربية وفقًا لهذه التفاصيل:
        - *الوصف*: {description}
        - *المجال*: {domain}
        - *الهدف من الإنشاء*: {goal}
        - *عدد الأعمدة*: {columns}
        - *عدد الصفوف*: {chunk_size}

        📌 *الإخراج يجب أن يكون JSON فقط، بدون أي شرح إضافي*:
        ```json
        {{
            "columns": ["اسم العمود 1", "اسم العمود 2", "اسم العمود 3", "اسم العمود 4", "اسم العمود 5"],
            "rows": [
                ["بيانات 1", "بيانات 2", "بيانات 3", "بيانات 4", "بيانات 5"],
                ["بيانات 1", "بيانات 2", "بيانات 3", "بيانات 4", "بيانات 5"]
            ]
        }}
        ```
        """

        retries = 0
        while retries < MAX_RETRIES:
            try:
                response = claude_client.messages.create(
                    model="claude-3-opus-20240229",
                    max_tokens=4096,
                    temperature=0.2,
                    messages=[{"role": "user", "content": prompt}]
                )

                raw_output = response.content[0].text

                # ✅ Extract only valid JSON (ignore extra text if present)
                json_match = re.search(r"\{.*\}", raw_output, re.DOTALL)
                if not json_match:
                    print(f"⚠️ Claude returned non-JSON response. Retrying... ({retries+1}/{MAX_RETRIES})")
                    retries += 1
                    continue

                dataset_chunk = json.loads(json_match.group().strip())  # ✅ Parse JSON safely

                if dataset_chunk.get("rows"):
                    if first_batch:
                        dataset["columns"] = dataset_chunk["columns"]
                        first_batch = False
                    else:
                        # ✅ Ensure column consistency
                        if dataset["columns"] != dataset_chunk["columns"]:
                            print(f"⚠️ Claude returned inconsistent columns. Fixing format...")
                            dataset_chunk = fix_columns(dataset_chunk, dataset["columns"])

                    all_rows.extend(dataset_chunk["rows"])
                    break  # ✅ Success, move to next batch

                print(f"⚠️ Claude returned invalid JSON. Retrying... ({retries+1}/{MAX_RETRIES})")
                retries += 1

            except json.JSONDecodeError:
                print(f"❌ Claude JSON Error: Invalid JSON format. Retrying... ({retries+1}/{MAX_RETRIES})")
                retries += 1

            except Exception as e:
                print(f"❌ Claude API Error: {e}")
                break  # No more retries if API call fails

    return {"columns": dataset["columns"], "rows": all_rows}


# ✅ **Helper Function to Fix Column Mismatch**
def fix_columns(dataset_chunk, correct_columns):
    """Ensures new batches match the columns of the first batch."""
    fixed_rows = []
    new_columns = dataset_chunk["columns"]

    for row in dataset_chunk["rows"]:
        new_row = []
        for col in correct_columns:
            if col in new_columns:
                new_row.append(row[new_columns.index(col)])
            else:
                new_row.append("")  # Fill missing columns with empty strings
        fixed_rows.append(new_row)

    return {"columns": correct_columns, "rows": fixed_rows}


import re
import json


def generate_gemini_dataset(description, domain, columns, rows, goal="غير محدد"):
    """Generates dataset in chunks using Gemini AI."""
    all_rows = []
    dataset = {"columns": [], "rows": []}
    BATCH_SIZE = 50  # Adjust batch size for stability
    first_batch = True

    for start_row in range(0, rows, BATCH_SIZE):
        chunk_size = min(BATCH_SIZE, rows - start_row)
        print(f"🔹 Requesting Gemini: {chunk_size} rows (from {start_row})")

        prompt = f"""
        أنت نموذج ذكاء اصطناعي. قم بإنشاء مجموعة بيانات باللغة العربية وفقًا لهذه التفاصيل:
        - الوصف: {description}
        - المجال: {domain}
        - الهدف من الإنشاء: {goal}
        - عدد الأعمدة: {columns}
        - عدد الصفوف: {chunk_size}

        📌 الإخراج يجب أن يكون JSON فقط، بدون أي شرح إضافي:
        json
        {{
            "columns": ["اسم العمود 1", "اسم العمود 2", "اسم العمود 3", "اسم العمود 4", "اسم العمود 5"],
            "rows": [
                ["بيانات 1", "بيانات 2", "بيانات 3", "بيانات 4", "بيانات 5"],
                ["بيانات 1", "بيانات 2", "بيانات 3", "بيانات 4", "بيانات 5"]
            ]
        }}

        """

        try:
            model = genai.GenerativeModel("gemini-1.5-pro-latest")
            response = model.generate_content(prompt)

            raw_output = response.text

            # ✅ Extract only valid JSON (ignore extra text if present)
            json_match = re.search(r"\{.*\}", raw_output, re.DOTALL)
            if not json_match:
                print(f"⚠️ Gemini returned non-JSON response. Skipping batch.")
                continue

            dataset_chunk = json.loads(json_match.group().strip())  # ✅ Parse JSON safely

            if dataset_chunk.get("rows"):
                if first_batch:
                    dataset["columns"] = dataset_chunk["columns"]
                    first_batch = False
                else:
                    # ✅ Ensure column consistency
                    if dataset["columns"] != dataset_chunk["columns"]:
                        dataset_chunk = fix_columns(dataset_chunk, dataset["columns"])

                all_rows.extend(dataset_chunk["rows"])

        except json.JSONDecodeError:
            print(f"❌ Gemini JSON Error: Invalid JSON format. Skipping batch.")
        except Exception as e:
            print(f"❌ Gemini API Error: {e}")

    return {"columns": dataset["columns"], "rows": all_rows}



@app.route('/generate_dataset', methods=['POST'])
def generate_dataset():
    """Generates dataset using GPT-4o, Claude 3, and Gemini, then saves them to MySQL."""
    try:
        data = request.json
        description = data.get("description", "لم يتم تقديم وصف")
        domain = data.get("domain", "عام")
        columns = int(data.get("columns", 5))
        rows = int(data.get("rows", 100))
        goal = data.get("goal", "غير محدد")

        print(f"🔹 Dataset Request: {description}, {domain}, {columns} cols, {rows} rows, Goal: {goal}")

        # ✅ Ensure user is logged in
        user_id = session.get('id')
        if not user_id:
            print("❌ Error: User not logged in.")
            return jsonify({"success": False, "error": "يجب تسجيل الدخول لحفظ البيانات."})

        # ✅ Generate datasets
        gpt4o_result = process_text_chunk(description, domain, columns, rows, goal)
        claude_result = generate_claude_dataset(description, domain, columns, rows, goal)
        gemini_result = generate_gemini_dataset(description, domain, columns, rows, goal)

        # ✅ Generate summaries
        gpt4o_summary = generate_summary(description, gpt4o_result) if gpt4o_result["columns"] else None
        claude_summary = generate_summary(description, claude_result) if claude_result["columns"] else None
        gemini_summary = generate_summary(description, gemini_result) if gemini_result["columns"] else None

        # ✅ Ensure at least one dataset is valid before inserting
        if not any([gpt4o_result["columns"], claude_result["columns"], gemini_result["columns"]]):
            print("❌ Error: No valid dataset generated.")
            return jsonify({"success": False, "error": "❌ لم يتم استخراج بيانات منظمة."})

        cursor = mysql.connection.cursor()

        # ✅ Save GPT-4o dataset
        gpt4o_id = None
        if gpt4o_result["columns"]:
            cursor.execute("""
                INSERT INTO tbl_datasets 
                (user_id, dataset_name, content_type, columns_data, rows_data, original_text, summary, model_used, goal) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                user_id, "GPT-4o Dataset", description,
                json.dumps(gpt4o_result["columns"], ensure_ascii=False),
                json.dumps(gpt4o_result["rows"], ensure_ascii=False),
                description, gpt4o_summary, "GPT-4o", goal
            ))
            gpt4o_id = cursor.lastrowid

        # ✅ Save Claude 3 dataset
        claude_id = None
        if claude_result["columns"]:
            cursor.execute("""
                INSERT INTO tbl_datasets 
                (user_id, dataset_name, content_type, columns_data, rows_data, original_text, summary, model_used, goal) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                user_id, "Claude 3 Dataset", description,
                json.dumps(claude_result["columns"], ensure_ascii=False),
                json.dumps(claude_result["rows"], ensure_ascii=False),
                description, claude_summary, "Claude-3", goal
            ))
            claude_id = cursor.lastrowid

        # ✅ Save Gemini dataset
        gemini_id = None
        if gemini_result["columns"]:
            cursor.execute("""
                    INSERT INTO tbl_datasets 
                    (user_id, dataset_name, content_type, columns_data, rows_data, original_text, summary, model_used, goal) 
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                user_id, "Gemini Dataset", description,
                json.dumps(gemini_result["columns"], ensure_ascii=False),
                json.dumps(gemini_result["rows"], ensure_ascii=False),
                description, gemini_summary, "Gemini", goal
            ))
            gemini_id = cursor.lastrowid

        cursor.execute("""
            INSERT INTO tbl_api_analysis 
            (user_id, original_text, goal, gpt4_datasets_id, claude_datasets_id, gemini_datasets_id) 
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (
            user_id,
            description,
            goal,
            gpt4o_id,
            claude_id,
            gemini_id
        ))

        mysql.connection.commit()
        cursor.close()

        return jsonify({
            "success": True,
            "gpt4o_id": gpt4o_id,
            "claude_id": claude_id,
            "gemini_id": gemini_id,
            "redirect_url": url_for("show_dataset_results",
                                    gpt4o_id=gpt4o_id, claude_id=claude_id, gemini_id=gemini_id)
        })

    except MySQLdb.Error as db_error:
        print(f"❌ MySQL Error in generate_dataset: {db_error}")
        return jsonify({"success": False, "error": "❌ خطأ في قاعدة البيانات."})

    except Exception as e:
        print(f"❌ Error in generate_dataset: {e}")
        return jsonify({"success": False, "error": str(e)})








import MySQLdb.cursors
import MySQLdb.cursors
from datetime import datetime

def convert_datetime(result):
    """Convert all datetime fields in a dictionary to strings before JSON serialization."""
    if result:
        for key in list(result.keys()):  # Ensure dictionary size remains unchanged during iteration
            if isinstance(result[key], datetime):
                result[key] = result[key].strftime('%Y-%m-%d %H:%M:%S')
    return result


@app.route('/dataset_results')
def show_dataset_results():
    """Retrieve and display dataset results for GPT-4o, Claude 3, and Gemini with pagination."""
    try:
        gpt4o_id = request.args.get("gpt4o_id")
        claude_id = request.args.get("claude_id")
        gemini_id = request.args.get("gemini_id")  # ✅ Added Gemini ID retrieval
        page = int(request.args.get("page", 1))
        rows_per_page = 40

        if not gpt4o_id and not claude_id and not gemini_id:
            return render_template("zero_api_result.html", error="⚠ لم يتم العثور على بيانات.")

        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)

        # ✅ Fetch GPT-4o dataset
        gpt4o_dataset = None
        if gpt4o_id:
            cursor.execute("SELECT * FROM tbl_datasets WHERE id = %s", (gpt4o_id,))
            gpt4o_dataset = cursor.fetchone()
            if gpt4o_dataset:
                gpt4o_dataset["columns"] = json.loads(gpt4o_dataset["columns_data"])
                gpt4o_dataset["rows"] = json.loads(gpt4o_dataset["rows_data"])

        # ✅ Fetch Claude 3 dataset
        claude_dataset = None
        if claude_id:
            cursor.execute("SELECT * FROM tbl_datasets WHERE id = %s", (claude_id,))
            claude_dataset = cursor.fetchone()
            if claude_dataset:
                claude_dataset["columns"] = json.loads(claude_dataset["columns_data"])
                claude_dataset["rows"] = json.loads(claude_dataset["rows_data"])

        # ✅ Fetch Gemini dataset
        gemini_dataset = None
        if gemini_id:
            cursor.execute("SELECT * FROM tbl_datasets WHERE id = %s", (gemini_id,))
            gemini_dataset = cursor.fetchone()
            if gemini_dataset:
                gemini_dataset["columns"] = json.loads(gemini_dataset["columns_data"])
                gemini_dataset["rows"] = json.loads(gemini_dataset["rows_data"])

        cursor.close()

        # ✅ Debugging output to confirm data
        print(f"✅ gpt4o_dataset: {gpt4o_dataset}")
        print(f"✅ claude_dataset: {claude_dataset}")
        print(f"✅ gemini_dataset: {gemini_dataset}")

        # ✅ Ensure at least one dataset exists before rendering
        if not gpt4o_dataset and not claude_dataset and not gemini_dataset:
            return render_template("zero_api_result.html", error="⚠ لم يتم العثور على بيانات.")

        return render_template("zero_api_result.html",
                               gpt4o_dataset=gpt4o_dataset,
                               claude_dataset=claude_dataset,
                               gemini_dataset=gemini_dataset,  # ✅ Pass Gemini dataset
                               page=page,
                               total_pages=(len(gpt4o_dataset["rows"]) + rows_per_page - 1) // rows_per_page if gpt4o_dataset else 1)

    except Exception as e:
        print(f"❌ Error in show_dataset_results: {e}")
        return render_template("zero_api_result.html", error="❌ حدث خطأ أثناء تحميل البيانات.")


def generate_summary(description, dataset):
    """Generate a concise summary based on user description and dataset."""
    try:
        # ✅ Ensure dataset is a dictionary, not a list
        if not isinstance(dataset, dict):
            print(f"❌ Invalid dataset type: {type(dataset)} - Expected dictionary")
            return "❌ تعذر توليد ملخص للبيانات بسبب تنسيق غير صحيح."

        # ✅ Ensure dataset contains valid keys
        if not dataset.get("columns") or not dataset.get("rows"):
            return "❌ لا يوجد بيانات كافية لتوليد ملخص."

        # ✅ Extract general insights
        num_columns = len(dataset["columns"])
        num_rows = len(dataset["rows"])
        general_insight = f"يعكس هذا الجدول بيانات متعلقة بـ {description} "

        return general_insight

    except Exception as e:
        print(f"❌ Error generating summary: {e}")
        return "❌ تعذر توليد ملخص للبيانات."











import os
import json
import MySQLdb
from flask import render_template, request

@app.route('/zero_api_result')
def show_dataset():
    """Fetch and display datasets from database or file for GPT-4o & Claude 3."""
    try:
        gpt4o_id = request.args.get("gpt4o_id")
        claude_id = request.args.get("claude_id")
        print(f"🔹 Received gpt4o_id: {gpt4o_id}, claude_id: {claude_id}")

        datasets = {}

        # ✅ 1. Load datasets from the database
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)

        if gpt4o_id:
            print(f"⏳ Fetching GPT-4o dataset with ID {gpt4o_id}...")
            cursor.execute("SELECT * FROM tbl_datasets WHERE id = %s", (gpt4o_id,))
            datasets["GPT-4o"] = cursor.fetchone()
            print("✅ GPT-4o dataset:", datasets["GPT-4o"])

        if claude_id:
            print(f"⏳ Fetching Claude 3 dataset with ID {claude_id}...")
            cursor.execute("SELECT * FROM tbl_datasets WHERE id = %s", (claude_id,))
            datasets["Claude 3"] = cursor.fetchone()
            print("✅ Claude 3 dataset:", datasets["Claude 3"])

        cursor.close()

        # ✅ 2. If no dataset found, return an error page
        if not datasets:
            print("❌ لم يتم العثور على بيانات في قاعدة البيانات.")
            return render_template("zero_api_result.html", error="❌ لم يتم العثور على بيانات.")

        # ✅ 3. Return datasets in the template
        return render_template("zero_api_result.html", datasets=datasets)

    except Exception as e:
        print(f"❌ خطأ في show_dataset: {e}")
        return render_template("zero_api_result.html", error="❌ حدث خطأ أثناء الإجراء.")





from openai import OpenAI


@app.route('/processing', methods=['GET', 'POST'])
def processing_options():
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    if 'uploaded_files' not in session:
        return redirect(url_for('upload_file'))

    if request.method == 'POST':
        # Retrieve processing option
        processing_option = request.form.get('processing')

        if not processing_option:
            # يرجع مع رسالة إذا ما اختار شيء
            return render_template('processing.html', message="يرجى اختيار طريقة المعالجة.")

        print(f"Processing option selected: {processing_option}")  # Debugging output

        apply_ner = 'apply_ner' in request.form
        apply_key_extraction = 'apply_key_extraction' in request.form
        top_n = int(request.form.get('top_n', 10))
        freq_order = request.form.get('freq_order', 'most')

        session['processing_option'] = processing_option
        session['apply_ner'] = apply_ner
        session['apply_key_extraction'] = apply_key_extraction
        session['top_n'] = top_n
        session['freq_order'] = freq_order

        return redirect(url_for('process_results'))

    return render_template('processing.html')



from time import sleep


@app.route('/process_results', methods=['GET'])
def process_results():
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    if 'uploaded_files' not in session or 'processing_option' not in session:
        return redirect(url_for('upload_file'))

    # Retrieve files and options from session
    sleep(3)  # Simulating backend processing delay

    uploaded_files = session.get('uploaded_files', [])
    processing_option = session.get('processing_option')
    print(f"Processing option in process_results: {processing_option}")
    apply_ner = session.get('apply_ner', False)
    apply_key_extraction = session.get('apply_key_extraction', False)
    top_n = session.get('top_n', 10)
    freq_order = session.get('freq_order', 'most')
    # Initialize merged text
    merged_text = ""

    # Load and merge text from the uploaded files
    for file_path in uploaded_files:
        file_extension = file_path.split('.')[-1].lower()

        # Handle .txt files
        if file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                merged_text += f.read() + "\n"

        # Handle .docx files
        elif file_extension == 'docx':
            document = Document(file_path)
            docx_text = "\n".join([paragraph.text for paragraph in document.paragraphs])
            merged_text += docx_text + "\n"

    # Clean and tokenize merged text
    cleaned_text = clean_arabic_text(merged_text)
    tokens = tokenize_text(cleaned_text)

    # Apply processing options
    if processing_option == "clean_preprocess":
        tokens = remove_stopwords(tokens)
    # txtword = " ".join(remove_stopwords(cleaned_text.split()))
    elif processing_option == "clean_stem":
        tokens = remove_stopwords(tokens)
        #  txtword = " ".join(remove_stopwords(cleaned_text.split()))

        tokens = stem_words(tokens)
    #  txtword = " ".join(stem_words(txtword.split()))

    # Generate word frequencies
    word_frequencies = generate_word_frequencies(tokens)
    sorted_word_frequencies = sorted(
        word_frequencies.items(),
        key=lambda x: x[1],
        reverse=True
    ) if freq_order == 'most' else sorted(
        word_frequencies.items(),
        key=lambda x: x[1]
    )
    top_words = sorted_word_frequencies[:top_n]

    # Generate visualizations
    unique_id = uuid.uuid4()
    wordcloud_filepath = os.path.join('static', f'wordcloud_{session["id"]}_{unique_id}.png')
    generate_wordcloud(' '.join(tokens), wordcloud_filepath)

    unigram_plot_filepath = os.path.join('static', f'unigram_plot_{session["id"]}_{unique_id}.png')
    plot_word_frequencies(top_words, f'Top {top_n} {freq_order.capitalize()} Words', unigram_plot_filepath)
    unigram_plot_url = url_for('static', filename=f'unigram_plot_{session["id"]}_{unique_id}.png')

    # Generate N-grams and their visualizations
    bigrams, trigrams = generate_ngrams(tokens)
    bigram_plot_filepath = os.path.join('static', f'bigram_plot_{session["id"]}_{unique_id}.png')
    trigram_plot_filepath = os.path.join('static', f'trigram_plot_{session["id"]}_{unique_id}.png')

    bigram_plot_url = None
    trigram_plot_url = None

    if bigrams:
        plot_ngrams(bigrams, 'Top Bigrams', bigram_plot_filepath)
        bigram_plot_url = url_for('static', filename=f'bigram_plot_{session["id"]}_{unique_id}.png')
    else:
        print("No bigrams found with enough frequency to plot.")

    if trigrams:
        plot_ngrams(trigrams, 'Top Trigrams', trigram_plot_filepath)
        trigram_plot_url = url_for('static', filename=f'trigram_plot_{session["id"]}_{unique_id}.png')
    else:
        print("No trigrams found with enough frequency to plot.")

    # Perform NER if selected
    ner_results = extract_ner(tokens, model_ner, tokenizer_ner) if apply_ner else None
    keyword_results = extract_keywords(tokens, kw_model) if apply_key_extraction else None

    # Save results in the database
    try:
        # Save results to database
        cursor = mysql.connection.cursor()
        sanitized_text = sanitize_text(merged_text)  # Sanitize the original text

        # Save results to database
        cursor = mysql.connection.cursor()
        cursor.execute('SET NAMES utf8mb4;')
        cursor.execute('SET CHARACTER SET utf8mb4;')
        cursor.execute('SET character_set_connection=utf8mb4;')

        query = """
            INSERT INTO tbl_results (
                user_id, filename, original_text, cleaned_text, word_frequencies, 
                wordcloud_path, unigram_plot_path, bigram_plot_path, trigram_plot_path, 
                created_at, processing_type, ner_results, keyword_results, original_filename
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), %s, %s, %s, %s)
        """

        word_frequencies_str = ','.join([f"{k}:{v}" for k, v in word_frequencies.items()])
        ner_results_json = json.dumps(ner_results)
        keyword_results_json = json.dumps(keyword_results)

        cursor.execute(query, (
            session['id'],
            # f"{unique_id}_{uploaded_files[0].split('/')[-1]}",  # Unique filename
            session.get('unique_file_name'),  # Unique filename
            merged_text,
            cleaned_text,
            word_frequencies_str,
            wordcloud_filepath,
            unigram_plot_filepath,
            bigram_plot_filepath,
            trigram_plot_filepath,
            processing_option,
            ner_results_json,
            keyword_results_json,
            session['original_filename'],
        ))

        mysql.connection.commit()

    except Exception as e:
        print(f"Error saving results to database: {e}")
        return render_template('processing.html', message="حدث خطأ أثناء حفظ النتائج.")
    finally:
        cursor.close()

    # Finalize and return the result once processing is complete
    return render_template(
        'result.html',
        original_text=merged_text,
        cleaned_text=cleaned_text,
        wordcloud_url=url_for('static', filename=f'wordcloud_{session["id"]}_{unique_id}.png'),
        unigram_plot_url=unigram_plot_url,
        bigram_plot_url=bigram_plot_url,
        trigram_plot_url=trigram_plot_url,
        word_frequencies=top_words,
        ner_results=ner_results if ner_results else ["لا يوجد"],
        keyword_results=keyword_results if keyword_results else ["لا يوجد"]
    )


def sanitize_text(text):
    """
    Remove unsupported characters by encoding to UTF-8 and ignoring errors.
    """
    return text.encode('utf-8', 'ignore').decode('utf-8')


# Function to generate a unique default name like "ملف 1", "ملف 2", etc.
def generate_unique_name(user_id):
    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cursor.execute('SELECT default_name FROM tbl_results WHERE user_id = %s', (user_id,))
    existing_names = {row['default_name'] for row in cursor.fetchall()}

    index = 1
    while f"ملف {index}" in existing_names:
        index += 1
    return f"ملف {index}"


from flask import Response, stream_with_context


@app.route('/update_file_name/<int:result_id>', methods=['POST'])
def update_file_name(result_id):
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    data = request.get_json()
    new_name = data.get('new_name')

    if new_name:
        try:
            cursor = mysql.connection.cursor()
            cursor.execute('UPDATE tbl_results SET default_name = %s WHERE id = %s AND user_id = %s',
                           (new_name, result_id, session['id']))
            mysql.connection.commit()
            return 'Success', 200
        except Exception as e:
            print(f"Error updating the file name: {e}")
            return 'Failed to update', 500
    return 'Invalid request', 400

@app.route('/delete_file/<int:result_id>/<result_type>', methods=['DELETE'])
def delete_file(result_id, result_type):
    if 'loggedin' not in session:
        return jsonify({"success": False, "message": "User not logged in"}), 401

    print(f"❌ خطأ في result_type: {result_type}")

    try:
        cursor = mysql.connection.cursor()

        if result_type == 'api_analysis':
            cursor.execute('DELETE FROM tbl_datasets WHERE id = %s AND user_id = %s',
                           (result_id, session['id']))
        if result_type == 'text_analysis':
            cursor.execute('DELETE FROM tbl_results WHERE id = %s AND user_id = %s',
                           (result_id, session['id']))

        if result_type == 'zero_shot':
            cursor.execute('DELETE FROM tbl_api_analysis WHERE id = %s AND user_id = %s',
                           (result_id, session['id']))

        if result_type == 'ai_verifications':
            cursor.execute('DELETE FROM tbl_data_verifications WHERE id = %s AND user_id = %s',
                           (result_id, session['id']))

        if result_type == 'human_verifications':
            cursor.execute('DELETE FROM tbl_data_verifications WHERE id = %s AND user_id = %s',
                           (result_id, session['id']))

        # if cursor.rowcount == 0:
        #     return jsonify({"success": False, "message": "File not found or unauthorized action"}), 404

        mysql.connection.commit()
        return jsonify({"success": True, "message": "File deleted successfully"}), 200

    except Exception as e:
        print(f"Error deleting file: {e}")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        cursor.close()

@app.route('/view_result/<int:result_id>', methods=['GET'])
def view_result(result_id):
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cursor.execute("""
        SELECT original_text, cleaned_text, word_frequencies, wordcloud_path, 
               unigram_plot_path, bigram_plot_path, trigram_plot_path, original_filename, 
               ner_results, keyword_results, filename
        FROM tbl_results
        WHERE id = %s AND user_id = %s
    """, (result_id, session['id']))

    result = cursor.fetchone()
    if result:
        result['word_frequencies'] = [
            tuple(item.split(':')) for item in result['word_frequencies'].split(',')
        ] if result['word_frequencies'] else []

        result['ner_results'] = json.loads(result['ner_results']) if result['ner_results'] and result['ner_results'] != 'null' else ["لا يوجد"]
        result['keyword_results'] = json.loads(result['keyword_results']) if result['keyword_results'] else ["لا يوجد"]
        result['filename_display'] = result['filename'].split('_', 1)[-1]  # Strip UUID from filename

        return render_template('view_result.html', result=result)
    else:
        return "Result not found."


# Full Arabic text cleaning function
def clean_arabic_text(text):
    text = mapper.map_string(text)  # Normalize using CamelTools CharMapper
    text = re.sub(r'ـ+', '', text)  # Remove Tatweel (ـــــ)
    text = re.sub(r'[ؗ-ًؚ-ْ]', '', text)  # Remove diacritics
    text = text.replace("ﻻ", "لا")  # Standardize ligatures
    text = re.sub(r'[أإآ]', 'ا', text)  # Normalize Hamzated Alif (أ, إ, آ) to bare Alif (ا)
    text = re.sub(r'ى', 'ي', text)  # Normalize Alif Maqsura (ى) to Ya (ي)
    text = re.sub(r'[^؀-ۿ\s]', '', text)  # Remove non-Arabic characters
    # Remove excess repetitions unless it's a valid double-letter word
    text = re.sub(r'(.)\1+', lambda m: m.group(0) if m.group(0) in valid_double_letter_words else m.group(1), text)
    text = re.sub(r'<[^>]+>', '', text)  # Remove HTML or special markup
    return text


# Tokenize text using Camel Tools
def tokenize_text(text):
    return simple_word_tokenize(text)


# Remove stopwords using NLTK's Arabic stopword list
def remove_stopwords(tokens):
    return [word for word in tokens if word not in arabic_stopwords]


# Perform stemming using NLTK's ISRIStemmer
def stem_words(tokens):
    return [stemmer.stem(word) for word in tokens]


# Generate word frequencies and return them
def generate_word_frequencies(tokens):
    word_counts = Counter(tokens)
    return word_counts


# Function to generate N-grams (bi-grams and tri-grams)
def generate_ngrams(tokens):
    bigrams = list(ngrams(tokens, 2))
    trigrams = list(ngrams(tokens, 3))

    bigram_freq = Counter(bigrams)
    trigram_freq = Counter(trigrams)

    # Filter out n-grams with frequency > 1
    frequent_bigrams = [(ngram, freq) for ngram, freq in bigram_freq.items() if freq > 1]
    frequent_trigrams = [(ngram, freq) for ngram, freq in trigram_freq.items() if freq > 1]

    # Format n-grams for display
    formatted_bigrams = [(' '.join(ngram), freq) for ngram, freq in frequent_bigrams]
    formatted_trigrams = [(' '.join(ngram), freq) for ngram, freq in frequent_trigrams]

    return formatted_bigrams, formatted_trigrams


# Plot word frequencies
def plot_word_frequencies(word_frequencies, title, filepath):
    reshaped_words = [get_display(arabic_reshaper.reshape(word)) for word, _ in word_frequencies]
    freqs = [count for _, count in word_frequencies]
    plt.figure(figsize=(10, 5))
    plt.bar(reshaped_words, freqs)
    plt.title(title, fontsize=16)
    plt.xticks(rotation=45, fontsize=14, ha='right')  # Align Arabic text properly
    plt.xlabel('Words', fontsize=14)
    plt.ylabel('Frequency', fontsize=14)
    plt.tight_layout()
    plt.savefig(filepath, bbox_inches='tight')
    plt.close()


# Plot N-grams
def plot_ngrams(ngrams_with_freq, title, filepath, top_n=20):
    # Sort n-grams by frequency in descending order and take the top N
    sorted_ngrams = sorted(ngrams_with_freq, key=lambda x: x[1], reverse=True)[:top_n]

    if not sorted_ngrams:
        print(f"No {title.lower()} found with enough frequency to plot.")
        return
    reshaped_ngrams = [get_display(arabic_reshaper.reshape(ngram)) for ngram, _ in sorted_ngrams]
    ngram_freqs = [freq for _, freq in sorted_ngrams]
    plt.figure(figsize=(10, 5))
    plt.bar(reshaped_ngrams, ngram_freqs)
    plt.title(title, fontsize=16)
    plt.xticks(rotation=45, fontsize=12, ha='right')  # Adjusted font size for readability
    plt.xlabel('N-grams', fontsize=14)
    plt.ylabel('Frequency', fontsize=14)
    plt.tight_layout()
    plt.savefig(filepath, bbox_inches='tight')
    plt.close()


import os
from wordcloud import WordCloud
from bidi.algorithm import get_display
import arabic_reshaper


# Function to generate a word cloud with the correct font path
def generate_wordcloud(text, filepath):
    # Construct the full path to 'Amiri-Regular.ttf'
    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Amiri-Regular.ttf')

    # Verify if the font file exists
    if not os.path.exists(font_path):
        raise FileNotFoundError(f"Font file not found at {font_path}")

    # Reshape and adjust text direction for Arabic
    reshaped_text = get_display(arabic_reshaper.reshape(text))

    # Generate the word cloud with the specified font
    wordcloud = WordCloud(
        font_path=font_path,
        background_color='white',
        width=800,
        height=400
    ).generate(reshaped_text)

    # Save the generated word cloud to the specified filepath
    wordcloud.to_file(filepath)


#  Named Entity Recognition (NER)
def extract_ner(text, model, tokenizer, start_token="▁"):
    if not isinstance(text, str):
        text = " ".join(text)  # Ensure text is a single string

    tokenized_sentence = tokenizer([text], padding=True, truncation=True, return_tensors="pt")
    tokenized_sentences = tokenized_sentence['input_ids'].numpy()
    with torch.no_grad():
        output = model(**tokenized_sentence)
    last_hidden_states = output[0].numpy()
    label_indices = np.argmax(last_hidden_states[0], axis=1)
    tokens = tokenizer.convert_ids_to_tokens(tokenized_sentences[0])
    special_tags = set(tokenizer.special_tokens_map.values())
    grouped_tokens = []
    for token, label_idx in zip(tokens, label_indices):
        if token not in special_tags:
            if not token.startswith(start_token) and len(token.replace(start_token, "").strip()) > 0:
                grouped_tokens[-1]["token"] += token
            else:
                grouped_tokens.append({"token": token, "label": custom_labels[label_idx]})
    ents = []
    prev_label = "O"
    for token in grouped_tokens:
        label = token["label"].replace("I-", "").replace("B-", "")
        if token["label"] != "O":
            if label != prev_label:
                ents.append({"token": [token["token"]], "label": label})
            else:
                ents[-1]["token"].append(token["token"])
        prev_label = label
    ents = [
        {"token": "".join(rec["token"]).replace(start_token, " ").strip(), "label": rec["label"]}
        for rec in ents
    ]
    return ents


# Keyword Extraction
def extract_keywords(text, kw_model):
    if not isinstance(text, str):
        text = " ".join(text)  # Ensure text is a single string
    keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words=None, top_n=10)
    return keywords


from datetime import datetime

@app.route('/profile')
def profile():
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)

    try:
        # Fetch user info
        cursor.execute('''
            SELECT id, username, email, first_name, second_name, 
                   major, birthdate, country, cv 
            FROM tbl_users 
            WHERE id = %s
        ''', (session['id'],))
        user = cursor.fetchone()

        # Fetch text analysis results
        cursor.execute('''
            SELECT 
                id,
                filename as display_filename,
                'text_analysis' as processing_type,
                created_at as formatted_date,
                wordcloud_path,
                unigram_plot_path,
                bigram_plot_path,
                original_filename,
                trigram_plot_path
            FROM tbl_results 
            WHERE user_id = %s
            AND original_text IS NOT NULL ORDER BY created_at DESC
        ''', (session['id'],))
        text_results = cursor.fetchall() or []

        # Fetch API analysis results
        cursor.execute('''
            SELECT 
                id,
                model_used,
                dataset_name as display_filename,
                'api_analysis' as processing_type,
                created_at as formatted_date,
                NULL as wordcloud_path,
                NULL as unigram_plot_path,
                NULL as bigram_plot_path,
                NULL as trigram_plot_path
            FROM tbl_datasets 
            WHERE user_id = %s
            And model_used IS NULL ORDER BY created_at DESC
        ''', (session['id'],))
        api_results = cursor.fetchall() or []

        # Fetch zero shot results
        cursor.execute('''
            SELECT 
            *
            FROM tbl_api_analysis
            WHERE user_id = %s
            ORDER BY created_at DESC
        ''', (session['id'],))
        zero_shot_results = cursor.fetchall() or []

        # Fetch ai verifications results
        cursor.execute('''
            SELECT 
            id,
            original_filename,
            created_at
            FROM tbl_data_verifications
            WHERE user_id = %s AND type = %s
            ORDER BY created_at DESC
        ''', (session['id'], 'ai_verification'))
        ai_verifications_results = cursor.fetchall() or []

        # Fetch human verifications results
        cursor.execute('''
            SELECT 
            id,
            original_filename,
            created_at
            FROM tbl_data_verifications
            WHERE user_id = %s AND type = %s
            ORDER BY created_at DESC
        ''', (session['id'], 'human_verification'))
        human_verifications_results = cursor.fetchall() or []

        # Combine and format results
        all_results = []
        formated_text_results = []
        formated_api_results = []
        formated_zero_shot_results = []

        # Process text results
        for result in text_results:
            result_dict = dict(result)
            if result_dict['formatted_date']:
                # Format the date as string
                result_dict['formatted_date'] = result_dict['formatted_date'].strftime('%Y-%m-%d %H:%M')
            formated_text_results.append(result_dict)

        # Process API results
        for result in api_results:
            result_dict = dict(result)
            if result_dict['formatted_date']:
                # Format the date as string
                result_dict['formatted_date'] = result_dict['formatted_date'].strftime('%Y-%m-%d %H:%M')
            formated_api_results.append(result_dict)

        # Process API results
        for result in zero_shot_results:
            result_dict = dict(result)
            if result_dict['created_at']:
                # Format the date as string
                result_dict['formatted_date'] = result_dict['created_at'].strftime('%Y-%m-%d %H:%M')
            formated_zero_shot_results.append(result_dict)

        # Sort results by date
        formated_text_results.sort(
            key=lambda x: x.get('formatted_date', ''),
            reverse=True
        )
        formated_api_results.sort(
            key=lambda x: x.get('formatted_date', ''),
            reverse=True
        )
        formated_zero_shot_results.sort(
            key=lambda x: x.get('formatted_date', ''),
            reverse=True
        )

        return render_template('profile.html',
                               user=user if user else {},
                               text_results=formated_text_results,
                               api_results=formated_api_results,
                               zero_shot_results=formated_zero_shot_results,
                               ai_verifications_results=ai_verifications_results,
                               human_verifications_results=human_verifications_results,
                               error_message=request.args.get('error_message'),
                               success_message=request.args.get('success_message'))

    except Exception as e:
        print(f"Error in profile route: {e}")
        return "حدث خطأ في تحميل الصفحة الشخصية", 500

    finally:
      cursor.close()

#
@app.route('/update_biography', methods=['POST'])
def update_biography():
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    user_id = session['id']
    # Get form data
    birthdate = request.form.get('birthdate', '').strip()
    country = request.form.get('country', '').strip()
    cv = request.form.get('cv', '').strip()

    # Validate the form data
    if not birthdate or not country or not cv:
        return redirect(url_for('profile', error_message="يرجى ملء جميع الحقول المطلوبة."))

    try:
        # Update the user information in the database
        cursor = mysql.connection.cursor()
        cursor.execute("""
            UPDATE tbl_users
            SET birthdate = %s, country = %s, cv = %s
            WHERE id = %s
        """, (birthdate, country, cv, user_id))
        mysql.connection.commit()
        cursor.close()

        return redirect(url_for('profile', success_message="تم تحديث السيرة الذاتية بنجاح."))
    except Exception as e:
        print(f"Error updating biography: {e}")
        return redirect(url_for('profile', error_message="حدث خطأ أثناء تحديث السيرة الذاتية."))


from flask import Flask, flash, redirect, render_template, request, session, url_for


@app.route('/update_user_info', methods=['POST'])
def update_user_info():
    if 'loggedin' not in session:
        return redirect(url_for('login'))

    user_id = session['id']
    username = request.form.get('username')
    email = request.form.get('email')
    first_name = request.form.get('first_name')
    second_name = request.form.get('second_name')
    major = request.form.get('major')  # Added the major field

    # Validate the inputs
    if not all([username, email, first_name, second_name, major]):
        return redirect(url_for('profile', error_message="يرجى ملء جميع الحقول المطلوبة."))

    try:
        cursor = mysql.connection.cursor()
        cursor.execute("""
            UPDATE tbl_users 
            SET username = %s, email = %s, first_name = %s, second_name = %s, major = %s
            WHERE id = %s
        """, (username, email, first_name, second_name, major, user_id))  # Updated to include the major field
        mysql.connection.commit()
        cursor.close()
    except Exception as e:
        print(f"Error updating user info: {e}")
        return redirect(url_for('profile', error_message="حدث خطأ أثناء تحديث المعلومات. يرجى المحاولة لاحقًا."))

    return redirect(url_for('profile', success_message="تم تحديث معلوماتك بنجاح."))


from flask import Flask, render_template, request, redirect, url_for, session, jsonify

import traceback



from flask import request, jsonify
import smtplib
from email.message import EmailMessage
import os




if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080,debug=True)
